from __future__ import annotations

import importlib.util
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\HK2_Nam3\Ltdd\DA_LTDD")
OUT = Path(r"D:\HK2_Nam3\Ltdd\Nhom11_TranNguyenBaoAn_SRS_CapNhat_2026.docx")
ASSET_DIR = ROOT / "build" / "srs_updated_assets"
LOGO = ROOT / "build" / "srs_media" / "image1.png"
TABLE_HELPER = Path(
    r"C:\Users\baoan\.codex\plugins\cache\openai-primary-runtime"
    r"\documents\26.601.10930\skills\documents\scripts\table_geometry.py"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
VERY_LIGHT = "F7F9FC"
MUTED = "666666"
WHITE = "FFFFFF"
BLACK = "111111"
GREEN = "2E7D32"
ORANGE = "A65E00"
RED = "9B1C1C"
PINK = "E0248A"
CONTENT_WIDTH = 9360
TABLE_INDENT = 120


spec = importlib.util.spec_from_file_location("table_geometry", TABLE_HELPER)
table_geometry = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(table_geometry)
apply_table_geometry = table_geometry.apply_table_geometry
column_widths_from_weights = table_geometry.column_widths_from_weights


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, *, bold=False, color=BLACK, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths, header=True, font_size=9.3):
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH,
        indent_dxa=TABLE_INDENT,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120},
    )
    table.style = "Table Grid"
    if header:
        set_repeat_table_header(table.rows[0])
        for cell in table.rows[0].cells:
            set_cell_shading(cell, LIGHT_BLUE)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=font_size, bold=True, color=NAVY)
    for row in table.rows[1 if header else 0 :]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_font(run, size=font_size, color=BLACK)


def add_table(doc, headers, rows, weights, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=NAVY, size=font_size)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(headers) > 2 else None
            set_cell_text(row.cells[i], value, size=font_size, align=align)
    style_table(
        table,
        column_widths_from_weights(weights, CONTENT_WIDTH),
        header=True,
        font_size=font_size,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def set_cell_border(cell, **edges):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in edges:
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edges[edge].items():
                element.set(qn(f"w:{key}"), str(value))


def add_field(paragraph, instruction, result=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = result
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, *, bold_label=None, note=False):
    p = doc.add_paragraph(style="Normal")
    if note:
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), VERY_LIGHT)
        pPr.append(shd)
    if bold_label and text.startswith(bold_label):
        first, rest = text.split(":", 1)
        r1 = p.add_run(first + ":")
        set_font(r1, bold=True, color=NAVY)
        r2 = p.add_run(rest)
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=9, italic=True, color=MUTED)
    return p


def set_picture_alt(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_major_break(doc):
    doc.add_page_break()


def add_status_callout(doc, title, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, VERY_LIGHT)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "18", "color": color},
        top={"val": "single", "sz": "4", "color": "D9E2F3"},
        bottom={"val": "single", "sz": "4", "color": "D9E2F3"},
        right={"val": "single", "sz": "4", "color": "D9E2F3"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_font(r2, size=10)
    apply_table_geometry(
        table,
        [CONTENT_WIDTH],
        table_width_dxa=CONTENT_WIDTH,
        indent_dxa=160,
        cell_margins_dxa={"top": 130, "bottom": 130, "start": 160, "end": 160},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def font(size=22, bold=False):
    path = r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"
    return ImageFont.truetype(path, size)


def rounded_box(draw, xy, fill, outline, radius=16, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color=(70, 90, 120), width=4):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    p1 = (x2 - ux * 14 + px * 7, y2 - uy * 14 + py * 7)
    p2 = (x2 - ux * 14 - px * 7, y2 - uy * 14 - py * 7)
    draw.polygon([end, p1, p2], fill=color)


def centered(draw, box, text, fnt, fill=(20, 30, 45)):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        heights.append(bbox[3] - bbox[1])
    total = sum(heights) + (len(lines) - 1) * 6
    y = y1 + (y2 - y1 - total) / 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=fnt)
        w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + 6


def make_architecture(path):
    img = Image.new("RGB", (1500, 800), "white")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "KIẾN TRÚC TRIỂN KHAI ỨNG DỤNG", font=font(34, True), fill=(23, 54, 93))
    boxes = {
        "ui": (70, 170, 390, 610),
        "logic": (480, 170, 850, 610),
        "firebase": (950, 120, 1405, 330),
        "local": (950, 410, 1405, 620),
    }
    rounded_box(d, boxes["ui"], (247, 249, 252), (46, 116, 181), 24, 4)
    rounded_box(d, boxes["logic"], (232, 238, 245), (31, 77, 120), 24, 4)
    rounded_box(d, boxes["firebase"], (255, 244, 250), (224, 36, 138), 24, 4)
    rounded_box(d, boxes["local"], (241, 248, 241), (46, 125, 50), 24, 4)
    centered(d, (90, 190, 370, 270), "Flutter Mobile UI", font(27, True))
    centered(d, (100, 280, 360, 570), "Đăng nhập & hồ sơ\nTrang chủ & điều hướng\nVí & giao dịch\nNgân sách & mục tiêu\nCông nợ & báo cáo\nTiện ích & thông báo", font(22))
    centered(d, (500, 190, 830, 270), "Controllers / Services", font(27, True))
    centered(d, (510, 280, 820, 570), "Xác thực dữ liệu\nTính số dư nguyên tử\nTính ngân sách\nTính tiến độ mục tiêu\nLọc và tổng hợp báo cáo\nTạo dữ liệu demo", font(22))
    centered(d, (970, 145, 1385, 305), "Firebase\nAuthentication\nCloud Firestore\nFirebase Storage", font(24, True))
    centered(d, (970, 435, 1385, 595), "Lưu trữ cục bộ\nSharedPreferences\nẢnh dự phòng trong\nDocuments Directory", font(24, True))
    arrow(d, (390, 390), (480, 390))
    arrow(d, (850, 300), (950, 225))
    arrow(d, (850, 490), (950, 515))
    d.text((55, 720), "Nền tảng hiện tại: Flutter 3.x, Dart 3.x, Firebase; giao diện tiếng Việt.", font=font(21), fill=(80, 80, 80))
    img.save(path)


def make_function_map(path):
    img = Image.new("RGB", (1500, 950), "white")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "BẢN ĐỒ CHỨC NĂNG AS-BUILT", font=font(34, True), fill=(23, 54, 93))
    user_box = (60, 350, 300, 600)
    rounded_box(d, user_box, (232, 238, 245), (31, 77, 120), 26, 4)
    centered(d, user_box, "NGƯỜI DÙNG\nĐÃ XÁC THỰC", font(26, True))
    modules = [
        ("Tài khoản", "Đăng ký, đăng nhập,\nGoogle, mật khẩu, hồ sơ"),
        ("Ví", "Khởi tạo, thêm, đổi tên,\nxóa, chi tiết, chuyển tiền"),
        ("Giao dịch", "Thu/chi, ảnh, tìm kiếm,\nlọc, sửa, xóa, định kỳ"),
        ("Kế hoạch", "Ngân sách, mục tiêu,\ngóp tiền, cảnh báo"),
        ("Công nợ", "Đi vay, cho vay,\ntất toán, nhắc hạn"),
        ("Phân tích", "Trang chủ, báo cáo tuần/\ntháng/năm, danh mục"),
        ("Tiện ích", "Thông báo 1/3/7 ngày,\ndữ liệu demo, hỗ trợ"),
    ]
    ys = [105, 220, 335, 450, 565, 680, 795]
    for (title, desc), y in zip(modules, ys):
        box = (500, y, 1380, y + 95)
        rounded_box(d, box, (247, 249, 252), (46, 116, 181), 18, 3)
        d.text((525, y + 14), title, font=font(23, True), fill=(31, 77, 120))
        d.text((760, y + 13), desc, font=font(20), fill=(35, 35, 35))
        arrow(d, (300, 475), (500, y + 47), width=3)
    img.save(path)


def make_data_model(path):
    img = Image.new("RGB", (1600, 1000), "white")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "MÔ HÌNH DỮ LIỆU FIRESTORE RÚT GỌN", font=font(34, True), fill=(23, 54, 93))
    boxes = {
        "users": (630, 120, 970, 245),
        "wallets": (160, 345, 500, 480),
        "tx": (80, 650, 520, 815),
        "categories": (620, 345, 980, 480),
        "budgets": (1080, 345, 1440, 480),
        "goals": (650, 650, 990, 815),
        "contrib": (1080, 650, 1500, 815),
        "recurring": (1050, 120, 1470, 245),
        "debts": (110, 120, 500, 245),
    }
    fills = [(232, 238, 245), (247, 249, 252), (255, 244, 250), (241, 248, 241)]
    labels = {
        "users": "users/{uid}\nhồ sơ người dùng",
        "wallets": "wallets/{walletId}\ntên, loại, số dư",
        "tx": "transactions/{txId}\nthu/chi/chuyển, số tiền,\ndanh mục, ngày, ảnh",
        "categories": "categories/{categoryId}\nthu/chi, biểu tượng, màu",
        "budgets": "budgets/{budgetId}\nhạn mức, kỳ, trạng thái",
        "goals": "goals/{goalId}\nmục tiêu, tiến độ, hạn",
        "contrib": "goal_contributions/{id}\nví, giao dịch, số tiền",
        "recurring": "recurring_transactions/{id}\nchu kỳ, hạn kế tiếp",
        "debts": "debts/{id}\nvay/cho vay, đã trả, hạn",
    }
    for idx, (key, box) in enumerate(boxes.items()):
        rounded_box(d, box, fills[idx % len(fills)], (46, 116, 181), 20, 3)
        centered(d, box, labels[key], font(21, key == "users"))
    arrow(d, (630, 180), (500, 395))
    arrow(d, (330, 480), (300, 650))
    arrow(d, (800, 245), (800, 345))
    arrow(d, (970, 180), (1050, 180))
    arrow(d, (970, 210), (1080, 395))
    arrow(d, (830, 480), (820, 650))
    arrow(d, (990, 735), (1080, 735))
    arrow(d, (630, 205), (500, 180))
    d.text((65, 900), "Ghi chú: giao dịch được lưu theo từng ví; các collection còn lại có trường userId để phân tách dữ liệu người dùng.", font=font(21), fill=(80, 80, 80))
    img.save(path)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section):
    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    first_header.paragraphs[0].clear()
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].clear()

    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("SRS - Ứng dụng quản lý thu chi cá nhân | Nhóm 11")
    set_font(r, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("Trang ")
    set_font(r1, size=8.5, color=MUTED)
    field = add_field(p, "PAGE", "1")
    set_font(field, size=8.5, color=MUTED)
    r2 = p.add_run(" | Phiên bản 2.0 - 13/06/2026")
    set_font(r2, size=8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    for text, size, bold in [
        ("BỘ CÔNG THƯƠNG", 12, True),
        ("TRƯỜNG ĐẠI HỌC CÔNG THƯƠNG TP. HỒ CHÍ MINH", 12, True),
        ("KHOA CÔNG NGHỆ THÔNG TIN", 12, True),
    ]:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=NAVY)
        r.add_break()
    if LOGO.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(10)
        lp.paragraph_format.space_after = Pt(14)
        logo = lp.add_run().add_picture(str(LOGO), width=Inches(1.35))
        set_picture_alt(logo, "Logo HUIT", "Logo Trường Đại học Công Thương Thành phố Hồ Chí Minh.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("ĐỒ ÁN LẬP TRÌNH DI ĐỘNG")
    set_font(r, size=15, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("ĐẶC TẢ YÊU CẦU PHẦN MỀM (SRS)")
    set_font(r, size=23, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    r = p2.add_run("ỨNG DỤNG QUẢN LÝ THU CHI CÁ NHÂN")
    set_font(r, size=19, bold=True, color=DARK_BLUE)

    add_status_callout(
        doc,
        "Bản cập nhật theo sản phẩm thực tế (as-built)",
        "Tài liệu phản ánh đầy đủ chức năng đang có trong project Flutter tại ngày 13/06/2026, thay thế các mô tả cũ về Admin, xuất dữ liệu và push notification.",
        color=PINK,
    )

    rows = [
        ("Giáo viên hướng dẫn", "Đinh Thị Tâm"),
        ("Nhóm thực hiện", "Nhóm 11"),
        ("Phạm Văn Tú", "2001230835"),
        ("Trần Nguyễn Bảo An", "2001230004"),
        ("Huỳnh Thị Bích Ngọc", "2001230573"),
        ("Phiên bản tài liệu", "2.0"),
    ]
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], label, bold=True, color=NAVY, size=10)
        set_cell_text(row.cells[1], value, size=10)
        set_cell_shading(row.cells[0], LIGHT_GRAY)
    apply_table_geometry(
        table,
        [2700, 6660],
        table_width_dxa=CONTENT_WIDTH,
        indent_dxa=TABLE_INDENT,
        cell_margins_dxa={"top": 85, "bottom": 85, "start": 120, "end": 120},
    )
    table.style = "Table Grid"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    r = p.add_run("TP. Hồ Chí Minh, tháng 6 năm 2026")
    set_font(r, size=11, italic=True, color=MUTED)
    doc.add_page_break()


def add_document_control(doc):
    add_heading(doc, "KIỂM SOÁT TÀI LIỆU", 1)
    add_table(
        doc,
        ["Mục", "Thông tin"],
        [
            ("Tên tài liệu", "Đặc tả yêu cầu phần mềm - Ứng dụng quản lý thu chi cá nhân"),
            ("Mã tài liệu", "QLTC-N11-SRS"),
            ("Phiên bản", "2.0"),
            ("Ngày cập nhật", "13/06/2026"),
            ("Trạng thái", "As-built / Phù hợp mã nguồn hiện tại"),
            ("Nền tảng", "Flutter, Dart, Firebase Authentication, Cloud Firestore, Firebase Storage"),
        ],
        [1.75, 4.75],
        font_size=10,
    )
    add_heading(doc, "Lịch sử phiên bản", 2)
    add_table(
        doc,
        ["Phiên bản", "Ngày", "Nội dung thay đổi", "Người thực hiện"],
        [
            ("1.0", "03/2026", "SRS ban đầu: khảo sát, yêu cầu và use case dự kiến.", "Nhóm 11"),
            (
                "2.0",
                "13/06/2026",
                "Đồng bộ toàn bộ chức năng đã triển khai; bổ sung quy tắc số dư, thông báo trong ứng dụng, dữ liệu demo, mô hình dữ liệu, tiêu chí nghiệm thu và truy vết.",
                "Nhóm 11",
            ),
        ],
        [0.8, 0.9, 3.9, 0.9],
        font_size=9.5,
    )
    add_heading(doc, "Các thay đổi phạm vi quan trọng", 2)
    add_table(
        doc,
        ["Nội dung SRS cũ", "Trạng thái trong sản phẩm hiện tại", "Cách phản ánh trong SRS 2.0"],
        [
            ("Cổng quản trị Admin", "Chưa có trong project Flutter hiện tại.", "Loại khỏi phạm vi phát hành 1.0; đưa vào định hướng tương lai."),
            ("Xuất Excel/CSV", "Đã loại bỏ khỏi giao diện theo yêu cầu.", "Không còn là yêu cầu chức năng."),
            ("Push notification nền", "Chưa tích hợp dịch vụ thông báo hệ điều hành.", "Đặc tả đúng là thông báo/nhắc nhở trong ứng dụng khi mở biểu tượng chuông."),
            ("Cho phép thay đổi email hồ sơ", "Email bị khóa khi chỉnh sửa hồ sơ.", "Bổ sung quy tắc BR-02."),
            ("Giao dịch có thể làm ví âm", "Đã chặn ở tầng dịch vụ bằng Firestore transaction.", "Bổ sung yêu cầu bắt buộc và tiêu chí nghiệm thu."),
        ],
        [1.6, 2.25, 2.65],
        font_size=9.2,
    )
    add_heading(doc, "MỤC LỤC", 1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "Mục lục sẽ được cập nhật khi mở tài liệu.")
    doc.add_page_break()


def add_chapter_1(doc):
    add_heading(doc, "1. GIỚI THIỆU", 1)
    add_heading(doc, "1.1. Mục đích tài liệu", 2)
    add_body(
        doc,
        "Tài liệu này đặc tả đầy đủ yêu cầu phần mềm của ứng dụng quản lý thu chi cá nhân QLTC_N11. Nội dung được xây dựng từ SRS ban đầu và kiểm chứng trực tiếp với cấu trúc mã nguồn, màn hình, controller, service và model hiện có trong project.",
    )
    add_body(
        doc,
        "SRS 2.0 là cơ sở cho nghiệm thu đồ án, kiểm thử chức năng, bảo trì và phát triển tiếp theo. Khi có khác biệt giữa ý tưởng ban đầu và sản phẩm đã triển khai, tài liệu ưu tiên mô tả hành vi đang tồn tại trong ứng dụng.",
    )
    add_heading(doc, "1.2. Phạm vi sản phẩm", 2)
    add_body(
        doc,
        "Ứng dụng hỗ trợ một người dùng cá nhân quản lý nhiều ví tiền, ghi nhận giao dịch thu/chi, chuyển tiền nội bộ, lập ngân sách, tạo mục tiêu tiết kiệm, theo dõi công nợ, xem báo cáo và nhận nhắc nhở trong ứng dụng.",
    )
    add_table(
        doc,
        ["Trong phạm vi phiên bản 1.0", "Ngoài phạm vi hiện tại"],
        [
            ("Ứng dụng Flutter cho thiết bị di động; giao diện tiếng Việt.", "Trang quản trị dành cho Admin."),
            ("Đăng nhập email/mật khẩu và Google.", "Xuất dữ liệu Excel, CSV hoặc PDF."),
            ("Lưu dữ liệu người dùng trên Firebase.", "Đồng bộ ngân hàng hoặc ví điện tử thực tế."),
            ("Nhắc công nợ và giao dịch định kỳ trong ứng dụng.", "Push notification nền khi ứng dụng đóng."),
            ("Tạo bộ dữ liệu demo trên tài khoản đang đăng nhập.", "Chia sẻ tài khoản hoặc ví giữa nhiều người."),
        ],
        [3.25, 3.25],
        font_size=9.5,
    )
    add_heading(doc, "1.3. Thuật ngữ và viết tắt", 2)
    add_table(
        doc,
        ["Thuật ngữ", "Giải thích"],
        [
            ("SRS", "Software Requirements Specification - Đặc tả yêu cầu phần mềm."),
            ("Ví", "Nguồn tiền độc lập như tiền mặt, ngân hàng, ví điện tử hoặc quỹ tiết kiệm."),
            ("Giao dịch", "Bản ghi thu nhập, chi tiêu hoặc luân chuyển tiền giữa các ví."),
            ("Ngân sách", "Hạn mức chi tiêu của một danh mục trong một khoảng thời gian."),
            ("Mục tiêu", "Khoản tiền cần tích lũy trước ngày mục tiêu."),
            ("Công nợ", "Khoản đi vay hoặc cho vay cần theo dõi thanh toán."),
            ("As-built", "Mô tả đúng trạng thái sản phẩm đã được triển khai."),
        ],
        [1.55, 4.95],
        font_size=9.7,
    )
    add_heading(doc, "1.4. Tài liệu tham chiếu", 2)
    add_body(doc, "Mã nguồn project: D:\\HK2_Nam3\\Ltdd\\DA_LTDD.")
    add_body(doc, "SRS gốc: Nhom11_TranNguyenBaoAn_SRS.docx, cập nhật lần cuối ngày 25/04/2026.")
    add_body(doc, "Thiết kế Figma của nhóm: https://bit.ly/4cl9mBZ.")


def add_chapter_2(doc):
    add_major_break(doc)
    add_heading(doc, "2. MÔ TẢ TỔNG QUAN", 1)
    add_heading(doc, "2.1. Bối cảnh và định hướng", 2)
    add_body(
        doc,
        "Sản phẩm hướng đến sinh viên và người đi làm có nhu cầu ghi chép tài chính cá nhân bằng điện thoại. Ứng dụng tập trung vào thao tác nhanh, dữ liệu trực quan và các ràng buộc giúp ngăn sai lệch số dư.",
    )
    add_heading(doc, "2.2. Tác nhân và hệ thống liên quan", 2)
    add_table(
        doc,
        ["Tác nhân/Hệ thống", "Vai trò"],
        [
            ("Người dùng", "Đăng ký, đăng nhập và quản lý toàn bộ dữ liệu tài chính của chính mình."),
            ("Firebase Authentication", "Xác thực email/mật khẩu, Google, đặt lại và cập nhật mật khẩu."),
            ("Cloud Firestore", "Lưu hồ sơ, ví, giao dịch, danh mục, ngân sách, mục tiêu, công nợ và định kỳ."),
            ("Firebase Storage", "Lưu ảnh chứng từ giao dịch khi tải lên thành công."),
            ("SharedPreferences", "Lưu lựa chọn ghi nhớ đăng nhập và cài đặt thông báo trên thiết bị."),
        ],
        [1.75, 4.75],
        font_size=9.5,
    )
    add_status_callout(
        doc,
        "Tác nhân hợp lệ của bản hiện tại",
        "Sản phẩm chỉ có tác nhân nghiệp vụ là Người dùng. Admin trong sơ đồ cũ chưa được triển khai và không được xem là tác nhân của phiên bản 1.0.",
        color=ORANGE,
    )
    add_heading(doc, "2.3. Kiến trúc sản phẩm", 2)
    arch = ASSET_DIR / "architecture.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    architecture = p.add_run().add_picture(str(arch), width=Inches(6.35))
    set_picture_alt(
        architecture,
        "Kiến trúc ứng dụng",
        "Sơ đồ kiến trúc Flutter kết nối Firebase Authentication, Cloud Firestore và Firebase Storage.",
    )
    add_caption(doc, "Hình 1. Kiến trúc triển khai của ứng dụng")
    add_heading(doc, "2.4. Bản đồ chức năng", 2)
    fmap = ASSET_DIR / "function_map.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    function_map = p.add_run().add_picture(str(fmap), width=Inches(6.35))
    set_picture_alt(
        function_map,
        "Bản đồ chức năng",
        "Bản đồ các nhóm chức năng tài khoản, ví, giao dịch, ngân sách, mục tiêu, công nợ, báo cáo và tiện ích.",
    )
    add_caption(doc, "Hình 2. Các nhóm chức năng đã triển khai")
    add_heading(doc, "2.5. Môi trường vận hành và công nghệ", 2)
    add_table(
        doc,
        ["Thành phần", "Công nghệ/phiên bản"],
        [
            ("Ứng dụng", "Flutter; Dart SDK ^3.10.8; Material UI."),
            ("Xác thực", "firebase_auth ^6.0.0; google_sign_in ^7.2.0."),
            ("Cơ sở dữ liệu", "cloud_firestore ^6.0.0."),
            ("Lưu ảnh", "firebase_storage ^13.4.0; image_picker ^1.2.2; path_provider ^2.1.5."),
            ("Biểu đồ", "fl_chart ^1.2.0."),
            ("Định dạng/ngôn ngữ", "intl ^0.20.2; locale vi_VN."),
            ("Lưu cục bộ", "shared_preferences ^2.5.5."),
            ("Phiên bản ứng dụng", "1.0.0+1 trong pubspec.yaml."),
        ],
        [1.8, 4.7],
        font_size=9.4,
    )


AUTH_REQS = [
    ("FR-AUTH-01", "Đăng ký bằng email", "Nhập họ tên, email, mật khẩu và xác nhận mật khẩu; tạo tài khoản Firebase và hồ sơ users/{uid}.", "Must"),
    ("FR-AUTH-02", "Kiểm tra dữ liệu đăng ký", "Họ tên 2-50 ký tự; email tối đa 100 ký tự; mật khẩu 8-32 ký tự, có chữ hoa, chữ thường, số và chỉ dùng a-z, A-Z, 0-9, @.", "Must"),
    ("FR-AUTH-03", "Đăng nhập email/mật khẩu", "Xác thực qua Firebase; hiển thị lỗi phù hợp khi sai thông tin, mất mạng hoặc gửi quá nhiều yêu cầu.", "Must"),
    ("FR-AUTH-04", "Đăng nhập Google", "Xác thực Google; tạo hồ sơ người dùng nếu chưa tồn tại.", "Must"),
    ("FR-AUTH-05", "Ghi nhớ đăng nhập", "Lưu lựa chọn trên thiết bị; nếu tắt, ứng dụng đăng xuất khi khởi động lại.", "Should"),
    ("FR-AUTH-06", "Khôi phục mật khẩu", "Gửi email đặt lại mật khẩu đến email hợp lệ đã nhập.", "Must"),
    ("FR-AUTH-07", "Tạo/chỉnh sửa hồ sơ", "Lưu họ tên, ngày sinh, email và ảnh đại diện; dữ liệu cũ được điền sẵn khi chỉnh sửa.", "Must"),
    ("FR-AUTH-08", "Khóa email hồ sơ", "Email chỉ đọc trong chế độ chỉnh sửa và hiển thị biểu tượng khóa.", "Must"),
    ("FR-AUTH-09", "Đổi/thiết lập mật khẩu", "Tài khoản mật khẩu phải xác thực lại; tài khoản nhà cung cấp ngoài nhận email thiết lập mật khẩu.", "Must"),
    ("FR-AUTH-10", "Đăng xuất", "Yêu cầu xác nhận trước khi đăng xuất Firebase.", "Must"),
]

WALLET_REQS = [
    ("FR-WAL-01", "Khởi tạo ví lần đầu", "Chọn loại ví, nhập số dư ban đầu không âm và đánh dấu hoàn tất onboarding.", "Must"),
    ("FR-WAL-02", "Danh sách ví", "Hiển thị tên, loại, số dư từng ví và tổng số dư.", "Must"),
    ("FR-WAL-03", "Thêm ví", "Tạo ví với tên, loại và số dư ban đầu; số tiền được định dạng theo vi_VN.", "Must"),
    ("FR-WAL-04", "Đổi tên ví", "Tên không để trống và tối đa 30 ký tự.", "Must"),
    ("FR-WAL-05", "Xóa ví", "Yêu cầu xác nhận; xóa ví và toàn bộ giao dịch con của ví.", "Must"),
    ("FR-WAL-06", "Xem chi tiết ví", "Hiển thị số dư và lịch sử giao dịch theo ví.", "Must"),
    ("FR-WAL-07", "Chuyển tiền nội bộ", "Chọn hai ví khác nhau, số tiền > 0, kiểm tra đủ số dư và cập nhật hai ví trong cùng Firestore transaction.", "Must"),
    ("FR-WAL-08", "Không cho phép ví âm", "Mọi cập nhật số dư trực tiếp hoặc gián tiếp phải bị từ chối nếu kết quả < 0.", "Must"),
]

TX_REQS = [
    ("FR-TXN-01", "Tạo giao dịch thu/chi", "Nhập loại, số tiền, danh mục, ví, ngày và ghi chú.", "Must"),
    ("FR-TXN-02", "Định dạng số tiền", "Ô nhập chỉ nhận chữ số và hiển thị dấu chấm hàng nghìn, ví dụ 100000 thành 100.000.", "Must"),
    ("FR-TXN-03", "Đính kèm ảnh", "Chọn ảnh từ thư viện; tải Firebase Storage, nếu lỗi thì lưu bản sao cục bộ.", "Should"),
    ("FR-TXN-04", "Cập nhật số dư nguyên tử", "Thu cộng tiền; chi trừ tiền; đọc số dư mới nhất và ghi giao dịch/số dư trong một Firestore transaction.", "Must"),
    ("FR-TXN-05", "Sửa giao dịch", "Cho phép đổi số tiền, loại, danh mục, ví, ngày, ghi chú và ảnh; hoàn tác ảnh hưởng cũ trước khi áp dụng dữ liệu mới.", "Must"),
    ("FR-TXN-06", "Xóa giao dịch", "Yêu cầu xác nhận, hoàn tác số dư và cập nhật liên kết ngân sách/mục tiêu.", "Must"),
    ("FR-TXN-07", "Chi tiết giao dịch", "Hiển thị danh mục, số tiền, loại, ngày, ví, ghi chú, trạng thái định kỳ và ảnh.", "Must"),
    ("FR-TXN-08", "Tìm kiếm không tải lại", "Lọc cục bộ theo ghi chú và danh mục; thao tác gõ không tạo lại truy vấn Firestore hoặc màn hình loading.", "Must"),
    ("FR-TXN-09", "Lọc giao dịch", "Lọc theo loại, tuần/tháng/năm, danh mục và ví; hỗ trợ xóa toàn bộ bộ lọc.", "Must"),
    ("FR-TXN-10", "Danh mục giao dịch", "Tạo, sửa, xóa danh mục thu/chi; cấu hình tên, biểu tượng và màu.", "Must"),
    ("FR-TXN-11", "Danh mục mặc định", "Khởi tạo các danh mục cơ bản cho người dùng khi cần.", "Should"),
    ("FR-TXN-12", "Đồng bộ ngân sách", "Tạo/sửa/xóa giao dịch chi phải tính lại ngân sách liên quan.", "Must"),
]

REC_REQS = [
    ("FR-REC-01", "Tạo thiết lập định kỳ", "Lưu tên, số tiền, loại, danh mục, ví, tần suất, ngày đến hạn, ngày kết thúc và ảnh tùy chọn.", "Must"),
    ("FR-REC-02", "Chu kỳ hỗ trợ", "Không lặp lại, hằng ngày, hằng tuần, hằng tháng và hằng năm.", "Must"),
    ("FR-REC-03", "Danh sách và trạng thái hạn", "Sắp xếp theo ngày kế tiếp; hiển thị còn bao nhiêu ngày hoặc quá hạn.", "Must"),
    ("FR-REC-04", "Sửa/xóa định kỳ", "Cho phép cập nhật hoặc dừng thiết lập sau khi xác nhận.", "Must"),
    ("FR-REC-05", "Thanh toán/thu định kỳ", "Khi xác nhận, tạo giao dịch thật, kiểm tra số dư với khoản chi và tăng ngày đến hạn kế tiếp.", "Must"),
    ("FR-REC-06", "Lịch sử định kỳ", "Hiển thị các giao dịch đã được tạo từ thiết lập định kỳ.", "Should"),
]

PLAN_REQS = [
    ("FR-BUD-01", "Tạo ngân sách", "Chọn danh mục, hạn mức > 0, kỳ ngày/tuần/tháng/năm/tùy chỉnh, ngày bắt đầu/kết thúc và ghi chú.", "Must"),
    ("FR-BUD-02", "Ngăn ngân sách chồng lấn", "Không cho tạo hai ngân sách cùng danh mục có khoảng thời gian giao nhau.", "Must"),
    ("FR-BUD-03", "Tính tiến độ ngân sách", "Tổng hợp giao dịch chi cùng danh mục và thời gian trên mọi ví.", "Must"),
    ("FR-BUD-04", "Trạng thái ngân sách", "SAFE < 80%; WARNING từ 80%; DANGER từ 90%; OVER_LIMIT khi > 100%.", "Must"),
    ("FR-BUD-05", "Quản lý ngân sách", "Xem danh sách/tổng quan/chi tiết; sửa, xóa và tính lại.", "Must"),
    ("FR-GOAL-01", "Tạo mục tiêu", "Nhập tên, số tiền > 0, ngày hoàn thành, ghi chú, màu và biểu tượng.", "Must"),
    ("FR-GOAL-02", "Nạp tiền vào mục tiêu", "Chọn ví, số tiền > 0; trừ ví và tạo giao dịch chi cùng contribution trong một transaction.", "Must"),
    ("FR-GOAL-03", "Tiến độ mục tiêu", "Tính current, remain, progress từ toàn bộ contribution.", "Must"),
    ("FR-GOAL-04", "Trạng thái mục tiêu", "ON_GOING; NEAR_TARGET từ 80%; COMPLETED từ 100%; FAILED nếu quá hạn và chưa đủ.", "Must"),
    ("FR-GOAL-05", "Quản lý mục tiêu", "Xem danh sách/tổng quan/chi tiết/lịch sử nạp; sửa và xóa.", "Must"),
    ("FR-GOAL-06", "Xóa mục tiêu không hoàn tiền", "Xóa mục tiêu và contribution; các giao dịch nạp đã phát sinh không được hoàn tự động.", "Must"),
]

DEBT_REPORT_REQS = [
    ("FR-DEBT-01", "Tạo công nợ", "Chọn đi vay/cho vay, người giao dịch, tổng tiền, đã thanh toán, ngày hạn và ghi chú.", "Must"),
    ("FR-DEBT-02", "Ràng buộc thanh toán", "Tổng tiền > 0; đã thanh toán >= 0 và không vượt tổng tiền.", "Must"),
    ("FR-DEBT-03", "Trạng thái công nợ", "OPEN; DUE_SOON khi còn tối đa 3 ngày; OVERDUE khi quá hạn; PAID khi còn lại bằng 0.", "Must"),
    ("FR-DEBT-04", "Quản lý công nợ", "Xem tổng hợp, sửa, xóa và tất toán sau xác nhận.", "Must"),
    ("FR-REP-01", "Tổng quan trang chủ", "Hiển thị thao tác nhanh, tổng thu/chi tháng, giao dịch gần đây, ngân sách và mục tiêu.", "Must"),
    ("FR-REP-02", "Báo cáo theo thời gian", "Tổng hợp theo tuần, tháng hoặc năm.", "Must"),
    ("FR-REP-03", "Loại báo cáo", "Thu nhập, chi tiêu hoặc chênh lệch.", "Must"),
    ("FR-REP-04", "Biểu đồ và phân rã", "Biểu đồ cột; nhóm dữ liệu theo danh mục; mở giao dịch chi tiết.", "Must"),
]

UTILITY_REQS = [
    ("FR-NOT-01", "Bật/tắt thông báo trong ứng dụng", "Lưu cấu hình cục bộ bằng SharedPreferences.", "Must"),
    ("FR-NOT-02", "Chọn loại nhắc", "Cho phép bật riêng công nợ đến hạn và giao dịch định kỳ.", "Must"),
    ("FR-NOT-03", "Thời gian nhắc trước", "Chọn 1, 3 hoặc 7 ngày.", "Must"),
    ("FR-NOT-04", "Xem thông báo", "Biểu tượng chuông tải khoản sắp đến hạn/quá hạn và hiển thị số tiền, loại, ngày, thời gian còn lại, ghi chú.", "Must"),
    ("FR-DEMO-01", "Tạo dữ liệu demo", "Sau xác nhận, tạo dữ liệu mẫu gồm hồ sơ, danh mục, ví, giao dịch, chuyển ví, ngân sách, mục tiêu, contribution, định kỳ và công nợ.", "Should"),
    ("FR-HELP-01", "Trợ giúp và hỗ trợ", "Hiển thị email, hotline và thời gian hỗ trợ mẫu.", "Could"),
]


def add_requirements_section(doc, title, reqs):
    add_heading(doc, title, 2)
    add_table(
        doc,
        ["Mã", "Chức năng", "Yêu cầu/Hành vi", "Ưu tiên"],
        reqs,
        [1.05, 1.65, 3.25, 0.55],
        font_size=8.8,
    )


def add_chapter_3(doc):
    add_major_break(doc)
    add_heading(doc, "3. YÊU CẦU CHỨC NĂNG", 1)
    add_body(
        doc,
        "Mỗi yêu cầu có mã định danh duy nhất để phục vụ kiểm thử và truy vết. Mức ưu tiên dùng MoSCoW: Must (bắt buộc), Should (nên có), Could (có thể có).",
        note=True,
    )
    add_requirements_section(doc, "3.1. Tài khoản, xác thực và hồ sơ", AUTH_REQS)
    add_requirements_section(doc, "3.2. Ví và nguồn tiền", WALLET_REQS)
    add_requirements_section(doc, "3.3. Giao dịch và danh mục", TX_REQS)
    add_requirements_section(doc, "3.4. Giao dịch định kỳ", REC_REQS)
    add_requirements_section(doc, "3.5. Ngân sách và mục tiêu tiết kiệm", PLAN_REQS)
    add_requirements_section(doc, "3.6. Công nợ và báo cáo", DEBT_REPORT_REQS)
    add_requirements_section(doc, "3.7. Thông báo, dữ liệu demo và hỗ trợ", UTILITY_REQS)


BUSINESS_RULES = [
    ("BR-01", "Số dư ví không âm", "Mọi giao dịch chi, chuyển tiền, thanh toán định kỳ, góp mục tiêu, sửa hoặc xóa giao dịch đều phải bảo đảm số dư kết quả >= 0."),
    ("BR-02", "Email hồ sơ bất biến", "Người dùng không được chỉnh sửa email trong màn hình hồ sơ; thay đổi danh tính phải xử lý qua cơ chế xác thực."),
    ("BR-03", "Tiền tệ vi_VN", "Ô nhập tiền hiển thị dấu chấm hàng nghìn; giá trị lưu trữ là số, không lưu ký tự phân cách."),
    ("BR-04", "Tính nguyên tử", "Các thao tác tác động đồng thời đến giao dịch và số dư sử dụng Firestore transaction để tránh trạng thái dở dang và số dư cũ."),
    ("BR-05", "Chuyển ví", "Ví nguồn và ví đích phải khác nhau; ví nguồn đủ tiền; tổng tài sản không đổi."),
    ("BR-06", "Ngân sách chồng lấn", "Một danh mục không có hai ngân sách có khoảng thời gian giao nhau."),
    ("BR-07", "Ngưỡng ngân sách", "80% cảnh báo; 90% nguy hiểm; vượt 100% là quá hạn mức."),
    ("BR-08", "Tiến độ mục tiêu", "Tổng contribution quyết định số đã tích lũy và trạng thái; không lấy số nhập thủ công."),
    ("BR-09", "Xóa mục tiêu", "Không tự động hoàn tiền đã góp; giao dịch chi đã phát sinh vẫn là lịch sử hợp lệ."),
    ("BR-10", "Công nợ", "Số đã thanh toán không vượt tổng tiền; tất toán đặt số đã trả bằng tổng tiền."),
    ("BR-11", "Nhắc nhở", "Nhắc nhở hiện chỉ được tính và hiển thị trong ứng dụng theo cấu hình 1/3/7 ngày."),
    ("BR-12", "Xóa ví", "Xóa ví đồng thời xóa toàn bộ giao dịch con; ứng dụng phải yêu cầu xác nhận rõ hậu quả."),
]


def add_chapter_4(doc):
    add_major_break(doc)
    add_heading(doc, "4. QUY TẮC NGHIỆP VỤ VÀ USE CASE", 1)
    add_heading(doc, "4.1. Quy tắc nghiệp vụ", 2)
    add_table(
        doc,
        ["Mã", "Tên quy tắc", "Nội dung"],
        BUSINESS_RULES,
        [0.9, 1.75, 3.85],
        font_size=9.0,
    )
    add_heading(doc, "4.2. Danh sách use case chính", 2)
    add_table(
        doc,
        ["Mã UC", "Tên use case", "Tác nhân", "Kết quả"],
        [
            ("UC-01", "Đăng ký/đăng nhập", "Người dùng", "Phiên xác thực hợp lệ và hồ sơ người dùng."),
            ("UC-02", "Khởi tạo và quản lý ví", "Người dùng", "Danh sách ví và số dư được cập nhật."),
            ("UC-03", "Ghi nhận giao dịch", "Người dùng", "Giao dịch được lưu, số dư và ngân sách đồng bộ."),
            ("UC-04", "Chuyển tiền giữa ví", "Người dùng", "Hai ví cập nhật nguyên tử, không âm."),
            ("UC-05", "Quản lý định kỳ", "Người dùng", "Thiết lập/lịch sử và ngày đến hạn kế tiếp."),
            ("UC-06", "Lập ngân sách", "Người dùng", "Hạn mức và trạng thái theo chi tiêu thực tế."),
            ("UC-07", "Tạo và nạp mục tiêu", "Người dùng", "Ví bị trừ, giao dịch và contribution được tạo."),
            ("UC-08", "Quản lý công nợ", "Người dùng", "Khoản vay/cho vay có trạng thái và số còn lại."),
            ("UC-09", "Xem báo cáo", "Người dùng", "Tổng hợp, biểu đồ và phân rã danh mục."),
            ("UC-10", "Xem nhắc nhở", "Người dùng", "Danh sách công nợ/định kỳ sắp hạn hoặc quá hạn."),
            ("UC-11", "Tạo dữ liệu demo", "Người dùng", "Tài khoản có bộ dữ liệu mẫu hoàn chỉnh."),
        ],
        [0.75, 1.65, 1.1, 3.0],
        font_size=9.0,
    )
    use_cases = [
        (
            "4.3. UC-03 - Ghi nhận giao dịch",
            "Tiền điều kiện: Người dùng đã đăng nhập, có ít nhất một ví và một danh mục.",
            "Luồng chính: Người dùng chọn Thu hoặc Chi, nhập số tiền, danh mục, ngày, ví, ghi chú và ảnh tùy chọn; hệ thống kiểm tra dữ liệu, đọc số dư mới nhất, tính số dư sau giao dịch, lưu giao dịch và cập nhật ví trong cùng transaction; sau đó tính lại ngân sách liên quan.",
            "Ngoại lệ: Thiếu trường bắt buộc, số tiền <= 0, ví không tồn tại hoặc giao dịch làm ví âm thì hệ thống từ chối và hiển thị lỗi.",
            "Hậu điều kiện: Số dư, lịch sử giao dịch, ngân sách và mục tiêu liên quan ở trạng thái nhất quán.",
        ),
        (
            "4.4. UC-04 - Chuyển tiền giữa ví",
            "Tiền điều kiện: Có ít nhất hai ví.",
            "Luồng chính: Chọn ví nguồn, ví đích, số tiền và ghi chú; hệ thống kiểm tra hai ví khác nhau và đủ số dư; cập nhật hai số dư và tạo hai bản ghi lịch sử trong một transaction.",
            "Ngoại lệ: Cùng ví, số tiền <= 0, ví không tồn tại hoặc số dư không đủ.",
            "Hậu điều kiện: Tổng tài sản không đổi; ví nguồn không âm.",
        ),
        (
            "4.5. UC-07 - Nạp tiền vào mục tiêu",
            "Tiền điều kiện: Mục tiêu đang tồn tại và có ví nguồn.",
            "Luồng chính: Nhập số tiền, chọn ví và ghi chú; hệ thống kiểm tra số dư; trừ ví, tạo giao dịch chi thuộc danh mục tiết kiệm và contribution; tính lại tiến độ mục tiêu.",
            "Ngoại lệ: Số tiền không hợp lệ hoặc vượt số dư ví.",
            "Hậu điều kiện: Contribution liên kết đúng giao dịch và mục tiêu.",
        ),
        (
            "4.6. UC-10 - Xem nhắc nhở",
            "Tiền điều kiện: Thông báo trong ứng dụng đang bật và có ít nhất một loại nhắc được chọn.",
            "Luồng chính: Người dùng nhấn chuông; hệ thống đọc cấu hình 1/3/7 ngày, tải công nợ và giao dịch định kỳ phù hợp, hiển thị loại, số tiền, ngày hạn, trạng thái thời gian và ghi chú.",
            "Ngoại lệ: Thông báo tắt, chưa chọn loại nhắc, không có dữ liệu hoặc lỗi mạng.",
            "Hậu điều kiện: Không thay đổi dữ liệu tài chính.",
        ),
    ]
    for title, pre, flow, exception, post in use_cases:
        add_heading(doc, title, 2)
        add_body(doc, pre, bold_label="Tiền điều kiện")
        add_body(doc, flow, bold_label="Luồng chính")
        add_body(doc, exception, bold_label="Ngoại lệ")
        add_body(doc, post, bold_label="Hậu điều kiện")


def add_chapter_5(doc):
    add_major_break(doc)
    add_heading(doc, "5. YÊU CẦU DỮ LIỆU", 1)
    add_heading(doc, "5.1. Mô hình dữ liệu tổng quát", 2)
    dmodel = ASSET_DIR / "data_model.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data_model = p.add_run().add_picture(str(dmodel), width=Inches(6.35))
    set_picture_alt(
        data_model,
        "Mô hình dữ liệu",
        "Sơ đồ các collection Firestore chính và quan hệ dữ liệu theo người dùng.",
    )
    add_caption(doc, "Hình 3. Các collection và quan hệ dữ liệu chính")
    add_heading(doc, "5.2. Từ điển dữ liệu", 2)
    add_table(
        doc,
        ["Collection/Đường dẫn", "Trường chính", "Quyền sở hữu/Quan hệ"],
        [
            ("users/{uid}", "fullName, email, birthday, phoneNumber, avatarUrl, loginProvider, onboardingCompleted", "Một hồ sơ cho mỗi Firebase uid."),
            ("users/{uid}/wallets/{walletId}", "name, type, balance, createdAt", "Ví thuộc người dùng."),
            (".../wallets/{walletId}/transactions/{txId}", "amount, type, categoryId, category, note, createdAt, imageUrl, isRecurring, walletId", "Giao dịch thuộc ví."),
            ("categories/{categoryId}", "userId, name, type, iconCode, colorHex, isDefault", "Lọc theo userId."),
            ("budgets/{budgetId}", "userId, categoryId, limitAmount, spentAmount, remainAmount, progressPercent, dates, status", "Liên kết danh mục và người dùng."),
            ("goals/{goalId}", "userId, targetAmount, currentAmount, remainAmount, progressPercent, targetDate, status", "Mục tiêu của người dùng."),
            ("goal_contributions/{id}", "goalId, walletId, transactionId, amount, note, createdAt", "Liên kết mục tiêu - ví - giao dịch."),
            ("recurring_transactions/{id}", "userId, name, amount, type, categoryId, walletId, frequency, nextDueDate, endDate", "Thiết lập định kỳ của người dùng."),
            ("debts/{id}", "userId, type, personName, amount, paidAmount, dueDate, note, status", "Công nợ của người dùng."),
        ],
        [2.05, 3.15, 1.3],
        font_size=8.7,
    )
    add_heading(doc, "5.3. Toàn vẹn và nhất quán", 2)
    add_body(doc, "Số dư ví và giao dịch phải được cập nhật nguyên tử đối với các luồng tạo, sửa, xóa, chuyển ví và góp mục tiêu.")
    add_body(doc, "Mọi truy vấn dữ liệu cấp cao phải lọc theo uid/userId của người dùng đang xác thực.")
    add_body(doc, "Các giá trị tính toán như spentAmount, progressPercent và remainAmount được tái tính từ dữ liệu nguồn thay vì tin cậy giá trị nhập tay.")
    add_body(doc, "Dữ liệu legacy không có walletId hoặc categoryId được xử lý bằng cơ chế tìm kiếm tương thích trong một số luồng.")
    add_status_callout(
        doc,
        "Yêu cầu triển khai Firebase Rules",
        "Repository hiện không cung cấp tệp rules để kiểm chứng. Khi phát hành, Firestore và Storage Rules phải giới hạn truy cập theo Firebase uid và không cho phép người dùng đọc/ghi dữ liệu của tài khoản khác.",
        color=RED,
    )


def add_chapter_6(doc):
    add_major_break(doc)
    add_heading(doc, "6. YÊU CẦU GIAO DIỆN VÀ TÍCH HỢP", 1)
    add_heading(doc, "6.1. Điều hướng chính", 2)
    add_table(
        doc,
        ["Tab", "Nội dung"],
        [
            ("Trang chủ", "Thao tác nhanh, tổng thu/chi, giao dịch gần đây, ngân sách và mục tiêu."),
            ("Ví", "Danh sách ví, tổng số dư, chi tiết ví và thêm ví."),
            ("Giao dịch", "Tất cả giao dịch, giao dịch thường, thiết lập định kỳ, tìm kiếm và lọc."),
            ("Biến động", "Báo cáo tuần/tháng/năm theo thu, chi và chênh lệch."),
            ("Tiện ích", "Ngân sách, mục tiêu, công nợ, hồ sơ, mật khẩu, thông báo, demo và hỗ trợ."),
        ],
        [1.25, 5.25],
        font_size=9.5,
    )
    add_heading(doc, "6.2. Quy tắc UI/UX bắt buộc", 2)
    add_table(
        doc,
        ["Mã", "Yêu cầu"],
        [
            ("UI-01", "Tất cả màn hình nhập liệu phải hiển thị trạng thái loading và vô hiệu hóa nút lưu khi đang xử lý."),
            ("UI-02", "Thao tác xóa, đăng xuất, tạo dữ liệu demo và tất toán phải có hộp thoại xác nhận."),
            ("UI-03", "Thông báo thành công/lỗi phải dùng SnackBar hoặc dialog rõ nghĩa bằng tiếng Việt."),
            ("UI-04", "Số tiền hiển thị với phân cách hàng nghìn và ký hiệu đồng; ô nhập tiền dùng dấu chấm theo vi_VN."),
            ("UI-05", "Danh sách trống phải có trạng thái rỗng và hành động tiếp theo phù hợp."),
            ("UI-06", "Tìm kiếm giao dịch phải phản hồi tại chỗ, không làm mất nội dung hoặc hiển thị loading sau mỗi ký tự."),
            ("UI-07", "Dữ liệu cũ phải được điền sẵn khi chỉnh sửa."),
            ("UI-08", "Email hồ sơ ở trạng thái khóa phải có biểu tượng và tooltip giải thích."),
        ],
        [0.9, 5.6],
        font_size=9.3,
    )
    add_heading(doc, "6.3. Giao diện tích hợp bên ngoài", 2)
    add_table(
        doc,
        ["Dịch vụ", "Dữ liệu trao đổi", "Xử lý lỗi"],
        [
            ("Firebase Authentication", "Email, mật khẩu, Google token, trạng thái phiên.", "Hiển thị mã lỗi thân thiện; yêu cầu đăng nhập lại khi cần."),
            ("Cloud Firestore", "Tài liệu và stream dữ liệu tài chính.", "Bắt exception, không cập nhật UI thành công khi commit thất bại."),
            ("Firebase Storage", "Byte ảnh chứng từ.", "Fallback lưu ảnh vào thư mục documents của ứng dụng."),
            ("Image Picker", "Đường dẫn ảnh trong thư viện.", "Cho phép bỏ qua ảnh."),
            ("SharedPreferences", "remember_me và cấu hình thông báo.", "Dùng giá trị mặc định nếu chưa có."),
        ],
        [1.65, 2.65, 2.2],
        font_size=9.2,
    )


NFRS = [
    ("NFR-01", "Hiệu năng", "Thao tác lọc/tìm kiếm trên dữ liệu đã tải phải phản hồi ngay; không phát sinh truy vấn mạng theo từng ký tự."),
    ("NFR-02", "Tính nhất quán", "Các thao tác thay đổi số dư phải dùng Firestore transaction và không để trạng thái ghi một phần."),
    ("NFR-03", "Bảo mật", "Xác thực bằng Firebase; dữ liệu phải được phân tách theo uid; mật khẩu không lưu trong Firestore."),
    ("NFR-04", "Khả dụng", "Hiển thị loading, lỗi mạng và trạng thái rỗng; không để ứng dụng crash khi dữ liệu thiếu."),
    ("NFR-05", "Dễ sử dụng", "Giao diện tiếng Việt, điều hướng 5 tab, biểu tượng quen thuộc, định dạng tiền và ngày vi_VN."),
    ("NFR-06", "Khả năng bảo trì", "Phân tách models, controllers, services, screens và widgets; mã yêu cầu truy vết đến module."),
    ("NFR-07", "Khả năng mở rộng", "Cấu trúc collection hỗ trợ nhiều ví, danh mục, ngân sách, mục tiêu và công nợ trên mỗi tài khoản."),
    ("NFR-08", "Tương thích", "Ứng dụng Flutter hướng đến Android; giao diện phải thích ứng kích thước màn hình điện thoại phổ biến."),
    ("NFR-09", "Địa phương hóa", "Mặc định vi_VN; hỗ trợ locale vi_VN và en_US ở cấp MaterialApp."),
    ("NFR-10", "Khôi phục lỗi", "Nếu tải ảnh Firebase thất bại, ứng dụng lưu cục bộ; nếu giao dịch thất bại, không thay đổi số dư."),
    ("NFR-11", "Kiểm thử", "Mã nguồn phải vượt flutter analyze, flutter test và build APK debug trước nghiệm thu."),
    ("NFR-12", "Riêng tư", "Không hiển thị dữ liệu tài chính của người dùng khác; không đưa thông tin nhạy cảm vào log hoặc dữ liệu demo công khai."),
]


def add_chapter_7(doc):
    add_major_break(doc)
    add_heading(doc, "7. YÊU CẦU PHI CHỨC NĂNG", 1)
    add_table(
        doc,
        ["Mã", "Nhóm", "Yêu cầu"],
        NFRS,
        [0.85, 1.35, 4.3],
        font_size=9.2,
    )
    add_heading(doc, "7.1. Tiêu chí chất lượng hiện tại", 2)
    add_body(
        doc,
        "Tại thời điểm cập nhật SRS, project đã vượt flutter analyze, bộ test định dạng tiền và build APK debug; ứng dụng đã được cài và mở thành công trên Android emulator API 33.",
    )
    add_body(
        doc,
        "Các yêu cầu hiệu năng định lượng 1-2 giây trong SRS cũ chưa có benchmark tự động. Phiên bản này chuyển chúng thành mục tiêu kiểm thử trên môi trường có mạng ổn định thay vì tuyên bố đã được chứng minh.",
    )


TESTS = [
    ("AT-01", "Đăng ký email hợp lệ", "Tài khoản và users/{uid} được tạo.", "FR-AUTH-01"),
    ("AT-02", "Mật khẩu thiếu chữ hoa/số", "Form từ chối và nêu quy tắc.", "FR-AUTH-02"),
    ("AT-03", "Quên mật khẩu với email hợp lệ", "Firebase gửi email đặt lại.", "FR-AUTH-06"),
    ("AT-04", "Sửa hồ sơ", "Thông tin cũ được điền; email không nhập được.", "FR-AUTH-07/08"),
    ("AT-05", "Nhập 100000 vào ô tiền", "Hiển thị 100.000, lưu giá trị 100000.", "FR-TXN-02"),
    ("AT-06", "Tạo chi lớn hơn số dư", "Bị từ chối; số dư và lịch sử không đổi.", "FR-WAL-08"),
    ("AT-07", "Sửa chi làm ví âm", "Bị từ chối nguyên tử.", "FR-TXN-05"),
    ("AT-08", "Xóa khoản thu khiến số dư âm", "Bị từ chối và giao dịch vẫn tồn tại.", "FR-TXN-06"),
    ("AT-09", "Chuyển tiền đủ số dư", "Hai ví cập nhật, tổng tài sản không đổi.", "FR-WAL-07"),
    ("AT-10", "Chuyển tiền thiếu số dư", "Bị từ chối; hai ví không đổi.", "FR-WAL-07/08"),
    ("AT-11", "Tìm kiếm bằng một từ", "Danh sách lọc tại chỗ, không xuất hiện loading.", "FR-TXN-08"),
    ("AT-12", "Lọc theo ví/danh mục/thời gian", "Chỉ hiển thị giao dịch phù hợp.", "FR-TXN-09"),
    ("AT-13", "Thanh toán định kỳ thiếu số dư", "Không tạo giao dịch và không tăng ngày hạn.", "FR-REC-05"),
    ("AT-14", "Tạo ngân sách chồng lấn", "Hệ thống từ chối.", "FR-BUD-02"),
    ("AT-15", "Chi đạt 80%, 90%, >100%", "Trạng thái lần lượt WARNING, DANGER, OVER_LIMIT.", "FR-BUD-04"),
    ("AT-16", "Góp mục tiêu vượt số dư", "Bị từ chối; không tạo contribution/giao dịch.", "FR-GOAL-02"),
    ("AT-17", "Mục tiêu đạt 100%", "Trạng thái COMPLETED.", "FR-GOAL-04"),
    ("AT-18", "Công nợ đã trả > tổng", "Form từ chối.", "FR-DEBT-02"),
    ("AT-19", "Công nợ còn 1 ngày", "Trạng thái DUE_SOON và xuất hiện khi nhắc trước 1 ngày.", "FR-DEBT-03/FR-NOT-03"),
    ("AT-20", "Tất toán công nợ", "paidAmount bằng amount, trạng thái PAID.", "FR-DEBT-04"),
    ("AT-21", "Báo cáo tuần/tháng/năm", "Tổng và cột biểu đồ đúng dữ liệu kỳ.", "FR-REP-02"),
    ("AT-22", "Tắt thông báo", "Nhấn chuông hiển thị thông báo đang tắt.", "FR-NOT-01"),
    ("AT-23", "Chọn nhắc trước 7 ngày", "Tải khoản đến hạn trong phạm vi 7 ngày.", "FR-NOT-03/04"),
    ("AT-24", "Tạo dữ liệu demo", "Các nhóm dữ liệu mẫu được tạo và hiển thị.", "FR-DEMO-01"),
    ("AT-25", "Xóa ví", "Có xác nhận; ví và giao dịch con bị xóa.", "FR-WAL-05"),
]


def add_chapter_8(doc):
    add_major_break(doc)
    add_heading(doc, "8. TIÊU CHÍ NGHIỆM THU VÀ KIỂM THỬ", 1)
    add_heading(doc, "8.1. Ma trận kiểm thử chấp nhận", 2)
    add_table(
        doc,
        ["Mã", "Kịch bản", "Kết quả mong đợi", "Truy vết"],
        TESTS,
        [0.75, 2.35, 2.75, 0.65],
        font_size=8.6,
    )
    add_heading(doc, "8.2. Kiểm thử biên bắt buộc", 2)
    add_table(
        doc,
        ["Miền dữ liệu", "Giá trị biên cần kiểm tra"],
        [
            ("Số tiền", "Rỗng; 0; 1; bằng số dư; lớn hơn số dư 1 đồng; số rất lớn."),
            ("Mật khẩu", "7/8/32/33 ký tự; thiếu chữ hoa, chữ thường hoặc số; chứa khoảng trắng/ký tự cấm."),
            ("Email", "Rỗng; sai định dạng; 100/101 ký tự; đã tồn tại."),
            ("Ngày", "Ngày bắt đầu bằng/kém/lớn hơn ngày kết thúc; ngày đến hạn hôm nay; quá hạn 1 ngày."),
            ("Ngân sách", "79.99%; 80%; 89.99%; 90%; 100%; lớn hơn 100%."),
            ("Công nợ", "Đã trả 0; bằng tổng; lớn hơn tổng 1 đồng."),
            ("Tìm kiếm", "Rỗng; 1 ký tự; một từ; không có kết quả; chữ hoa/thường."),
        ],
        [1.7, 4.8],
        font_size=9.4,
    )
    add_heading(doc, "8.3. Điều kiện phát hành", 2)
    for release_condition in [
        "Không có lỗi crash hoặc màn hình đỏ trong các luồng chính.",
        "flutter analyze không có issue; flutter test đạt; flutter build apk --debug thành công.",
        "Tất cả kịch bản Must trong ma trận nghiệm thu đạt trên ít nhất một thiết bị Android hoặc emulator.",
        "Firebase Rules được rà soát và triển khai đúng quyền sở hữu dữ liệu trước khi phát hành công khai.",
    ]:
        p = add_body(doc, release_condition)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            set_font(run, size=9.5)


def add_chapter_9(doc):
    heading = add_heading(doc, "9. TRUY VẾT YÊU CẦU VÀ MÃ NGUỒN", 1)
    heading.paragraph_format.page_break_before = True
    add_table(
        doc,
        ["Nhóm yêu cầu", "Module chính"],
        [
            ("FR-AUTH-*", "lib/views/auth; lib/main.dart; Firebase Authentication."),
            ("FR-WAL-*", "lib/views/wallet/screens; lib/views/wallet/services/wallet_service.dart."),
            ("FR-TXN-*", "add_transaction_screen.dart; transaction_menu_screen.dart; transaction_service.dart; category_controller.dart."),
            ("FR-REC-*", "add_recurring_transaction_screen.dart; recurring_transaction_controller.dart; transaction_menu_screen.dart."),
            ("FR-BUD-*", "lib/views/budget; lib/controllers/budget_controller.dart."),
            ("FR-GOAL-*", "lib/views/goal; lib/controllers/goal_controller.dart."),
            ("FR-DEBT-*", "lib/views/debt; lib/controllers/debt_controller.dart."),
            ("FR-REP-*", "home_screen.dart; report_screen.dart; các widget tổng quan."),
            ("FR-NOT-*", "notification_settings_screen.dart; notification_settings_service.dart; notification_reminder_content.dart."),
            ("FR-DEMO-*", "lib/services/demo_data_service.dart; utility_screen.dart."),
        ],
        [1.45, 5.05],
        font_size=9.3,
    )
    add_heading(doc, "9.1. Cấu trúc dữ liệu demo", 2)
    add_body(
        doc,
        "Bộ dữ liệu mẫu hiện tạo bốn ví, nhiều giao dịch thu/chi và chuyển ví, năm ngân sách, bốn mục tiêu cùng contribution, năm thiết lập định kỳ và ba khoản công nợ. Dữ liệu bao gồm cả trạng thái an toàn, sắp đạt, hoàn thành, quá hạn và vượt mức để trình diễn đầy đủ UI.",
    )
    add_heading(doc, "9.2. Hạn chế và định hướng mở rộng", 2)
    add_table(
        doc,
        ["Hạn chế hiện tại", "Đề xuất mở rộng"],
        [
            ("Nhắc nhở chỉ xuất hiện khi người dùng mở ứng dụng và nhấn chuông.", "Tích hợp flutter_local_notifications hoặc FCM, lập lịch nền và xin quyền thông báo."),
            ("Chưa có Firebase Rules trong repository để kiểm chứng.", "Version-control rules, emulator test và kiểm thử truy cập chéo uid."),
            ("Ảnh hồ sơ có thể lưu đường dẫn cục bộ.", "Đồng nhất tải ảnh hồ sơ lên Firebase Storage và dọn ảnh cũ."),
            ("Xóa mục tiêu không hoàn tiền.", "Cung cấp tùy chọn hoàn tiền có transaction nguyên tử và lịch sử rõ ràng."),
            ("Chưa xuất dữ liệu theo quyết định sản phẩm hiện tại.", "Chỉ khôi phục khi có yêu cầu; ưu tiên Android Share Sheet và quyền riêng tư."),
            ("Báo cáo chưa có benchmark và so sánh kỳ trước đang tắt.", "Tối ưu truy vấn, cache tổng hợp và bật so sánh kỳ trước."),
            ("Chưa có kiểm thử service Firebase tự động.", "Dùng Firebase Emulator Suite và unit/integration test cho số dư, ngân sách, mục tiêu."),
        ],
        [3.1, 3.4],
        font_size=9.2,
    )
    add_status_callout(
        doc,
        "Kết luận phạm vi",
        "Phiên bản hiện tại đã vượt phạm vi quản lý thu/chi cơ bản ban đầu nhờ có Google Sign-In, ngân sách nhiều kỳ, mục tiêu gắn giao dịch, công nợ có trạng thái, nhắc nhở cấu hình, dữ liệu demo và các bảo vệ toàn vẹn số dư.",
        color=GREEN,
    )


def add_appendix(doc):
    add_major_break(doc)
    add_heading(doc, "PHỤ LỤC A. PHÂN CÔNG NHÓM", 1)
    add_table(
        doc,
        ["Họ và tên", "MSSV", "Phạm vi phụ trách ban đầu"],
        [
            ("Huỳnh Thị Bích Ngọc", "2001230573", "Giao diện quản lý hệ thống & người dùng; quản lý ví tiền."),
            ("Trần Nguyễn Bảo An", "2001230004", "Quản lý danh mục & giao dịch; nhóm tiện ích mở rộng."),
            ("Phạm Văn Tú", "2001230835", "Ngân sách & mục tiêu; quản lý vay nợ."),
        ],
        [2.0, 1.3, 3.2],
        font_size=9.5,
    )
    add_heading(doc, "PHỤ LỤC B. CHECKLIST ĐỐI CHIẾU SẢN PHẨM", 1)
    add_table(
        doc,
        ["Hạng mục", "Trạng thái SRS 2.0"],
        [
            ("Tài khoản, Google, mật khẩu, hồ sơ", "Đã đặc tả"),
            ("Ví, chuyển tiền, chặn số dư âm", "Đã đặc tả"),
            ("Giao dịch, ảnh, tìm kiếm, lọc, định dạng tiền", "Đã đặc tả"),
            ("Danh mục và giao dịch định kỳ", "Đã đặc tả"),
            ("Ngân sách và mục tiêu", "Đã đặc tả"),
            ("Công nợ và nhắc hạn", "Đã đặc tả"),
            ("Trang chủ và báo cáo", "Đã đặc tả"),
            ("Thông báo trong ứng dụng", "Đã đặc tả đúng phạm vi"),
            ("Dữ liệu demo", "Đã đặc tả"),
            ("Admin và xuất file", "Ghi rõ ngoài phạm vi"),
        ],
        [3.9, 2.6],
        font_size=9.5,
    )


def create_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    make_architecture(ASSET_DIR / "architecture.png")
    make_function_map(ASSET_DIR / "function_map.png")
    make_data_model(ASSET_DIR / "data_model.png")

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        configure_header_footer(section)

    doc.core_properties.title = "SRS 2.0 - Ứng dụng quản lý thu chi cá nhân"
    doc.core_properties.subject = "Đặc tả yêu cầu phần mềm cập nhật theo project"
    doc.core_properties.author = "Nhóm 11"
    doc.core_properties.comments = "Cập nhật ngày 13/06/2026 dựa trên mã nguồn as-built."

    add_cover(doc)
    add_document_control(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    add_chapter_5(doc)
    add_chapter_6(doc)
    add_chapter_7(doc)
    add_chapter_8(doc)
    add_chapter_9(doc)
    add_appendix(doc)

    for section in doc.sections:
        configure_section(section)
        configure_header_footer(section)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    create_document()
